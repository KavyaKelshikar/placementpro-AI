import os
import re
from pypdf import PdfReader
from docx import Document

def extract_text_from_pdf(file_path):
    """Extract text from PDF using pypdf."""
    try:
        reader = PdfReader(file_path)
        text = ""
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        return text
    except Exception as e:
        print(f"Error reading PDF: {e}")
        return ""

def extract_text_from_docx(file_path):
    """Extract text from Word Document using python-docx."""
    try:
        doc = Document(file_path)
        text = []
        for paragraph in doc.paragraphs:
            text.append(paragraph.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text.append(cell.text)
        return "\n".join(text)
    except Exception as e:
        print(f"Error reading DOCX: {e}")
        return ""

def extract_text(file_path):
    """Determine file extension and extract text."""
    _, ext = os.path.splitext(file_path.lower())
    if ext == '.pdf':
        return extract_text_from_pdf(file_path)
    elif ext in ['.docx', '.doc']:
        return extract_text_from_docx(file_path)
    elif ext in ['.txt', '.rtf']:
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
        except Exception as e:
            print(f"Error reading TXT: {e}")
            return ""
    else:
        return ""

def calculate_resume_score_and_suggestions(text):
    """
    Calculate a resume score out of 100 and suggest improvements.
    Categorized score weights:
    - Experience Section: 25 pts (and up to 5 pts bonus for action verbs)
    - Education Section: 20 pts
    - Projects Section: 20 pts
    - Skills Section: 20 pts
    - Contact info: 15 pts (email 5 pts, phone 5 pts, social links 5 pts)
    """
    score = 0
    suggestions = []
    text_lower = text.lower()

    # 1. Contact Details (Max 15 pts)
    has_email = bool(re.search(r'[\w\.-]+@[\w\.-]+\.\w+', text))
    has_phone = bool(re.search(r'\+?\d{3,4}[\s-]?\d{3,4}[\s-]?\d{4}|\b\d{10}\b', text))
    has_links = any(domain in text_lower for domain in ['github.com', 'linkedin.com', 'portfolio'])

    if has_email:
        score += 5
    else:
        suggestions.append("Add a professional email address to the header.")

    if has_phone:
        score += 5
    else:
        suggestions.append("Provide a phone number so recruiters can easily contact you.")

    if has_links:
        score += 5
    else:
        suggestions.append("Link your GitHub, LinkedIn, or personal portfolio profiles to showcase your work and network.")

    # 2. Work Experience (Max 25 pts)
    has_experience = any(kw in text_lower for kw in ['experience', 'work history', 'intern', 'employment', 'job history'])
    if has_experience:
        score += 25
        # Check for action verbs to ensure quality description (bonus up to 5 pts)
        action_verbs = ['developed', 'designed', 'implemented', 'led', 'managed', 'created', 'built', 'optimized', 'engineered', 'increased']
        matched_verbs = [verb for verb in action_verbs if verb in text_lower]
        if len(matched_verbs) >= 3:
            # Experience quality is good
            pass
        else:
            suggestions.append("Strengthen your descriptions using action verbs (e.g. 'designed', 'implemented', 'optimized') to illustrate achievements.")
    else:
        suggestions.append("Include a 'Work Experience' or 'Internships' section to outline your professional background.")

    # 3. Education (Max 20 pts)
    has_education = any(kw in text_lower for kw in ['education', 'academic', 'university', 'college', 'degree', 'cgpa', 'gpa'])
    if has_education:
        score += 20
    else:
        suggestions.append("Include an 'Education' section detailing your degree, institution, and major fields of study.")

    # 4. Projects (Max 20 pts)
    has_projects = any(kw in text_lower for kw in ['projects', 'personal projects', 'academic projects', 'portfolio projects'])
    if has_projects:
        score += 20
    else:
        suggestions.append("Add a 'Projects' section describing applications you built to demonstrate practical coding/engineering experience.")

    # 5. Skills (Max 20 pts)
    has_skills = any(kw in text_lower for kw in ['skills', 'technologies', 'technical skills', 'core competencies'])
    if has_skills:
        score += 20
    else:
        suggestions.append("Add a clear 'Skills' section listing languages, frameworks, databases, and tools you know.")

    # 6. Overall Length and Density Check
    word_count = len(text.split())
    if word_count < 100:
        suggestions.append("Your resume content seems brief. Elaborate on your projects, courses, and skills to provide depth.")
        score = max(0, score - 15)
    elif word_count > 1000:
        suggestions.append("Your resume is very long. Condense descriptions to fit key achievements, aiming for 1-2 pages maximum.")
        score = max(0, score - 5)

    return {
        "score": min(100, score),
        "suggestions": suggestions,
        "wordCount": word_count
    }
